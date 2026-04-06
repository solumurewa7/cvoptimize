# services/resume_parser.py
#
# Handles extracting plain text from uploaded resume files.
#
# We support two file formats:
#   - PDF  → parsed with PyMuPDF (imported as `fitz`)
#   - DOCX → parsed with python-docx
#
# WHY extract text instead of storing the original file?
# - Cheaper: no file storage service needed (no S3, no Supabase Storage)
# - Simpler: all our AI/NLP analysis just reads a plain string
# - Safer: no risk of serving malicious file uploads back to users
#
# The tradeoff: we can't show the user their original formatted resume.
# For CVOptimize's use case (analysis only), that's fine.

import io
import fitz          # PyMuPDF — pip install pymupdf
from docx import Document  # python-docx


class UnsupportedFileTypeError(Exception):
    """Raised when the uploaded file is not a PDF or DOCX."""
    pass


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """
    Extract plain text from a PDF or DOCX file.

    Args:
        file_bytes: The raw bytes of the uploaded file.
        file_type:  "pdf" or "docx".

    Returns:
        A single string containing all the text in the document.

    Raises:
        UnsupportedFileTypeError: if file_type is not "pdf" or "docx".
    """
    if file_type == "pdf":
        return _extract_from_pdf(file_bytes)
    elif file_type == "docx":
        return _extract_from_docx(file_bytes)
    else:
        raise UnsupportedFileTypeError(f"Unsupported file type: {file_type}")


def _extract_from_pdf(file_bytes: bytes) -> str:
    """
    Use PyMuPDF to extract text from a PDF.

    fitz.open() can accept raw bytes via a stream.
    We iterate over every page and concatenate the text.
    """
    # fitz.open() with stream= reads from bytes instead of a file path.
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []
    for page in doc:
        # get_text() returns the page's plain text.
        # "text" mode preserves line breaks; we strip trailing whitespace.
        pages.append(page.get_text("text").strip())
    doc.close()

    # Join pages with double newline so sections stay visually separated.
    return "\n\n".join(pages)


def _extract_from_docx(file_bytes: bytes) -> str:
    """
    Use python-docx to extract text from a DOCX file.

    python-docx reads a file-like object, so we wrap the bytes in io.BytesIO.
    We iterate over every paragraph and join them.
    """
    # io.BytesIO wraps raw bytes in a file-like interface.
    # python-docx's Document() reads from it just like a real file.
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    return "\n".join(paragraphs)


def count_words(text: str) -> int:
    """Count the number of words in extracted text."""
    return len(text.split())
